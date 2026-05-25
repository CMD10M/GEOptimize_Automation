import base64
import json
import os
import re
import tempfile
from datetime import datetime
from pathlib import Path
from docx.shared import Inches

import streamlit as st
import time
from anthropic._exceptions import (
    OverloadedError,
    RateLimitError,
    APIConnectionError,
)
from anthropic import Anthropic
from docx import Document
from docxtpl import DocxTemplate
from dotenv import load_dotenv
from docx.shared import Pt

from report_prompt import SYSTEM_PROMPT, USER_PROMPT, REQUIRED_KEYS
from proposal_prompt import PROPOSAL_SYSTEM_PROMPT, PROPOSAL_USER_PROMPT

load_dotenv()

BASE_DIR = Path(__file__).resolve().parent

FEASIBILITY_TEMPLATE_PATH = (
    BASE_DIR
    / "templates"
    / "GEOptimize_GHX_Feasibility_Autofill_Template.docx"
)

PROPOSAL_TEMPLATE_PATH = (
    BASE_DIR
    / "templates"
    / "GEOptimize_Proposal_Master_Template.docx"
)

CAD_TEMPLATE_PATH = (
    BASE_DIR
    / "templates"
    / "Vertical_GHX.docx"
)

SLINKY_CAD_TEMPLATE_PATH = (
    BASE_DIR
    / "templates"
    / "Slinky_GHX.docx"
)

HDD_CAD_TEMPLATE_PATH = (
    BASE_DIR
    / "templates"
    / "HDD_GHX.docx"
)



PROPOSAL_DATABASE_DIR = BASE_DIR / "proposal_database"
PROPOSAL_DATABASE_DIR.mkdir(exist_ok=True)


def pdf_block_from_bytes(data: bytes):
    encoded = base64.standard_b64encode(data).decode("utf-8")
    return {
        "type": "document",
        "source": {
            "type": "base64",
            "media_type": "application/pdf",
            "data": encoded,
        },
    }


def pdf_block(uploaded_file):
    return pdf_block_from_bytes(uploaded_file.getvalue())


def text_block(label: str, text: str):
    return {"type": "text", "text": f"\n\n--- {label} ---\n{text}"}


def extract_docx_text(data: bytes) -> str:
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(data)
        tmp_path = tmp.name
    try:
        doc = Document(tmp_path)
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    finally:
        try:
            os.remove(tmp_path)
        except OSError:
            pass


def uploaded_file_to_content_block(uploaded_file, label: str):
    suffix = Path(uploaded_file.name).suffix.lower()
    data = uploaded_file.getvalue()
    if suffix == ".pdf":
        return pdf_block_from_bytes(data)
    if suffix == ".docx":
        return text_block(label, extract_docx_text(data))
    if suffix in [".txt", ".md"]:
        return text_block(label, data.decode("utf-8", errors="replace"))
    return text_block(label, f"Unsupported file type for direct extraction: {uploaded_file.name}")


def database_file_to_content_block(path: Path):
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return pdf_block_from_bytes(path.read_bytes())
    if suffix == ".docx":
        return text_block(f"Existing proposal example: {path.name}", extract_docx_text(path.read_bytes()))
    if suffix in [".txt", ".md"]:
        return text_block(f"Existing proposal example: {path.name}", path.read_text(encoding="utf-8", errors="replace"))
    return None


def extract_text(message) -> str:
    parts = []
    for block in message.content:
        if getattr(block, "type", None) == "text":
            parts.append(block.text)
    return "\n".join(parts).strip()


def parse_json_response(text: str) -> dict:
    cleaned = text.strip()
    fenced = re.search(r"```(?:json)?\s*(\{.*?\})\s*```", cleaned, flags=re.DOTALL)
    if fenced:
        cleaned = fenced.group(1)
    else:
        start = cleaned.find("{")
        end = cleaned.rfind("}")
        if start != -1 and end != -1 and end > start:
            cleaned = cleaned[start : end + 1]
    return json.loads(cleaned)


def validate_context(context: dict) -> dict:
    for key in REQUIRED_KEYS:
        context.setdefault(key, "Not provided")
    return context


def render_feasibility_docx(context: dict) -> bytes:
    with tempfile.TemporaryDirectory() as tmpdir:
        out_path = Path(tmpdir) / "GHX_Feasibility_Report.docx"
        doc = DocxTemplate(str(FEASIBILITY_TEMPLATE_PATH))
        doc.render(context)
        doc.save(str(out_path))
        return out_path.read_bytes()

def render_cad_request_docx(context: dict, rough_sketch_image=None) -> bytes:
    with tempfile.TemporaryDirectory() as tmpdir:
        out_path = Path(tmpdir) / "Vertical_GHX_CAD_Request_Form.docx"

        doc = DocxTemplate(str(CAD_TEMPLATE_PATH))
        doc.render(context)
        doc.save(str(out_path))

        final_doc = Document(str(out_path))
        replace_marker_with_image(
            final_doc,
            "[[ROUGH_SKETCH_IMAGE]]",
            rough_sketch_image,
            width_inches=6.0,
        )
        final_doc.save(str(out_path))

        return out_path.read_bytes()


def remove_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)


def as_list(value):
    if isinstance(value, list):
        return value

    if isinstance(value, str):
        return [
            line.strip("•- ").strip()
            for line in value.splitlines()
            if line.strip()
        ]

    return []


def replace_marker_with_scope(doc: Document, marker: str, scope_sections):
    if not isinstance(scope_sections, list):
        scope_sections = []

    for paragraph in list(doc.paragraphs):
        if marker in paragraph.text:
            for section in scope_sections:
                heading = section.get("heading", "").strip()
                intro = section.get("intro", "").strip()
                tasks = section.get("tasks", [])

                if heading:
                    p = paragraph.insert_paragraph_before(heading)
                    p.style = "Heading 2"

                if intro:
                    p = paragraph.insert_paragraph_before(intro)
                    p.style = "Normal"
                for task in tasks:
                    p = paragraph.insert_paragraph_before(str(task).strip())
                    try:
                        p.style = "GEOptimize Bullet"
                    except KeyError:
                        p.style = "Normal"
                        p.text = "• " + p.text

            remove_paragraph(paragraph)
            break

def replace_marker_with_image(doc: Document, marker: str, uploaded_image, width_inches: float = 6.0):
    if not uploaded_image:
        return

    with tempfile.NamedTemporaryFile(
        suffix=Path(uploaded_image.name).suffix,
        delete=False,
    ) as tmp:
        tmp.write(uploaded_image.getvalue())
        image_path = tmp.name

    try:
        for paragraph in list(doc.paragraphs):
            if marker in paragraph.text:
                paragraph.text = ""
                run = paragraph.add_run()
                run.add_picture(image_path, width=Inches(width_inches))
                break
    finally:
        try:
            os.remove(image_path)
        except OSError:
            pass


def replace_marker_with_bullets(doc: Document, marker: str, items):
    items = as_list(items)

    for paragraph in list(doc.paragraphs):
        if marker in paragraph.text:
            for item in items:
                p = paragraph.insert_paragraph_before(str(item).strip())
                try:
                    p.style = "GEOptimize Bullet"
                except KeyError:
                    p.style = "Normal"
                    p.text = "• " + p.text

            remove_paragraph(paragraph)
            break

def replace_marker_with_terms(doc: Document, marker: str, terms_text: str):
    for paragraph in list(doc.paragraphs):
        if marker in paragraph.text:

            for line in terms_text.splitlines():
                clean = line.strip()

                if not clean:
                    paragraph.insert_paragraph_before("")

                elif re.match(r"^\d+\.\s", clean):
                    p = paragraph.insert_paragraph_before(clean)
                    p.style = "Legal"

                elif re.match(r"^\([a-z]\)", clean):
                    p = paragraph.insert_paragraph_before(clean)
                    p.style = "Legal"

                else:
                    p = paragraph.insert_paragraph_before(clean)
                    p.style = "Legal"

            remove_paragraph(paragraph)
            break


def render_proposal_docx(context: dict, terms_region: str) -> bytes:
    with tempfile.TemporaryDirectory() as tmpdir:
        rendered_path = Path(tmpdir) / "proposal_rendered.docx"
        final_path = Path(tmpdir) / "GEOptimize_Proposal_Draft.docx"

        doc = DocxTemplate(str(PROPOSAL_TEMPLATE_PATH))
        doc.render(context)
        doc.save(str(rendered_path))

        final_doc = Document(str(rendered_path))

        replace_marker_with_scope(
            final_doc,
            "[[SCOPE_OF_WORK]]",
            context.get("scope_sections", [])
        )

        replace_marker_with_bullets(
            final_doc,
            "[[INFORMATION_REQUIRED]]",
            context.get("information_required", [])
        )

        replace_marker_with_bullets(
            final_doc,
            "[[DELIVERABLES]]",
            context.get("deliverables", [])
        )

        if terms_region == "USA":
            terms_path = BASE_DIR / "templates" / "terms" / "us_terms.txt"
        else:
            terms_path = BASE_DIR / "templates" / "terms" / "canada_terms.txt"

        terms_text = terms_path.read_text(encoding="utf-8")

        replace_marker_with_terms(
            final_doc,
            "[[TERMS_AND_CONDITIONS]]",
            terms_text
        )

        final_doc.save(str(final_path))
        return final_path.read_bytes()

def render_slinky_cad_docx(context: dict, rough_sketch_image=None) -> bytes:
    with tempfile.TemporaryDirectory() as tmpdir:
        out_path = Path(tmpdir) / "Slinky_GHX_CAD_Request_Form.docx"

        doc = DocxTemplate(str(SLINKY_CAD_TEMPLATE_PATH))
        doc.render(context)
        doc.save(str(out_path))

        final_doc = Document(str(out_path))
        replace_marker_with_image(
            final_doc,
            "[[ROUGH_SKETCH_IMAGE]]",
            rough_sketch_image,
            width_inches=6.0,
        )
        final_doc.save(str(out_path))

        return out_path.read_bytes()

def render_hdd_cad_docx(context: dict, rough_sketch_image=None) -> bytes:
    with tempfile.TemporaryDirectory() as tmpdir:
        out_path = Path(tmpdir) / "HDD_GHX_CAD_Request_Form.docx"

        doc = DocxTemplate(str(HDD_CAD_TEMPLATE_PATH))
        doc.render(context)
        doc.save(str(out_path))

        final_doc = Document(str(out_path))
        replace_marker_with_image(
            final_doc,
            "[[ROUGH_SKETCH_IMAGE]]",
            rough_sketch_image,
            width_inches=6.0,
        )
        final_doc.save(str(out_path))

        return out_path.read_bytes()

def call_claude(client: Anthropic, model: str, max_tokens: int, system_prompt: str, content: list):
    retries = 3

    for attempt in range(retries):
        try:
            return client.messages.create(
                model=model,
                max_tokens=int(max_tokens),
                system=system_prompt,
                messages=[{"role": "user", "content": content}],
            )

        except Exception as exc:
            error_text = str(exc)

            overloaded = (
                "529" in error_text
                or "Overloaded" in error_text
                or "rate_limit" in error_text
            )

            if not overloaded or attempt == retries - 1:
                raise exc

            wait_seconds = 5 * (attempt + 1)

            st.warning(
                f"Claude API busy. Retrying in {wait_seconds} seconds..."
            )

            time.sleep(wait_seconds)


st.set_page_config(page_title="GEOptimize Automation", layout="centered")
st.title("GEOptimize Automation")

with st.sidebar:
    st.header("Claude Settings")
    api_key = st.text_input(
        "Anthropic API key",
        value=os.getenv("ANTHROPIC_API_KEY", ""),
        type="password",
        help="You can also place this in a .env file as ANTHROPIC_API_KEY.",
    )
    model = st.text_input(
        "Claude model",
        value=os.getenv("CLAUDE_MODEL", "claude-opus-4-7"),
        help="Use the model ID shown in your Claude Console. For cost savings, use a Sonnet model if preferred.",
    )
    max_tokens = st.number_input("Max output tokens", min_value=2000, max_value=30000, value=10000, step=500)
    page = st.radio(
        "Tool",
        [
            "Feasibility Report Generator",
            "Proposal Generator",
            "Vertical GHX CAD Request Form",
            "Slinky GHX CAD Request Form",
            "HDD GHX CAD Request Form",
        ],
    )

if page == "Feasibility Report Generator":
    st.header("GHX Feasibility Report Generator")
    st.caption("Upload the proposal and GLD output PDF, then generate a formatted Word report.")

    proposal_pdf = st.file_uploader("Proposal PDF", type=["pdf"], key="feasibility_proposal")
    gld_pdf = st.file_uploader("GLD output PDF", type=["pdf"], key="feasibility_gld")

    if st.button("Generate Feasibility Report", type="primary"):
        if not api_key:
            st.error("Enter your Anthropic API key in the sidebar or .env file.")
            st.stop()
        if not proposal_pdf or not gld_pdf:
            st.error("Upload both the proposal PDF and the GLD output PDF.")
            st.stop()

        client = Anthropic(api_key=api_key)
        content = [pdf_block(proposal_pdf), pdf_block(gld_pdf), {"type": "text", "text": USER_PROMPT}]

        with st.status("Analyzing documents with Claude and generating report...", expanded=True) as status:
            st.write("Sending PDFs to Claude...")
            message = call_claude(client, model, max_tokens, SYSTEM_PROMPT, content)
            st.write("Parsing JSON response...")
            raw_text = extract_text(message)
            try:
                context = validate_context(parse_json_response(raw_text))
                context["report_date"] = datetime.now().strftime("%B %d, %Y")
            except Exception as exc:
                st.error("Claude returned output that could not be parsed as JSON.")
                st.text_area("Raw Claude output", raw_text, height=400)
                st.exception(exc)
                st.stop()

            st.write("Filling Word template...")
            docx_bytes = render_feasibility_docx(context)
            status.update(label="Report generated", state="complete")

        st.success("Done. Download the completed Word report below.")
        st.download_button(
            label="Download GHX Feasibility Report (.docx)",
            data=docx_bytes,
            file_name="GHX_Feasibility_Report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        with st.expander("View generated JSON"):
            st.json(context)

elif page == "Proposal Generator":
    st.header("Proposal Generator")
    st.caption("Use existing proposal examples, uploaded project documents, and notes to draft a new proposal.")

    st.subheader("1. Existing proposal database")
    st.write("Place reusable proposal examples in the local `proposal_database` folder. Supported formats: PDF, DOCX, TXT, MD.")
    database_files = sorted([
        p for p in PROPOSAL_DATABASE_DIR.iterdir()
        if p.suffix.lower() in [".pdf", ".docx", ".txt", ".md"]
    ])
    if database_files:
        selected_database_files = st.multiselect(
            "Select proposal examples to use as style/scope references",
            options=[p.name for p in database_files],
            default=[],
        )
    else:
        selected_database_files = []
        st.info("No proposal examples found yet. Add files to the `proposal_database` folder, or upload examples below.")

    uploaded_examples = st.file_uploader(
        "Optional: upload additional example proposals",
        type=["pdf", "docx", "txt", "md"],
        accept_multiple_files=True,
        key="uploaded_examples",
    )

    st.subheader("2. New project documents")
    project_docs = st.file_uploader(
        "Upload site plans, RFPs, client notes, building documents, or related project files",
        type=["pdf", "docx", "txt", "md"],
        accept_multiple_files=True,
        key="project_docs",
    )

    st.subheader("3. Project notes")
    project_notes = st.text_area(
        "Enter any additional project information, scope direction, pricing notes, deadlines, exclusions, or assumptions",
        height=220,
        placeholder="Example: Client is evaluating preliminary GHX sizing only. Include energy model review as optional scope. Pricing to be left as placeholder.",
    )
    st.subheader("4. Terms & Conditions")

    terms_region = st.radio(
        "Select Terms & Conditions",
        ["Canada", "USA"],
        horizontal=True,
    )

    if st.button("Generate Proposal Draft", type="primary"):
        if not api_key:
            st.error("Enter your Anthropic API key in the sidebar or .env file.")
            st.stop()
        if not project_docs and not project_notes.strip():
            st.error("Upload at least one new project document or enter project notes.")
            st.stop()

        client = Anthropic(api_key=api_key)
        content = []

        selected_paths = [PROPOSAL_DATABASE_DIR / name for name in selected_database_files]
        for path in selected_paths:
            block = database_file_to_content_block(path)
            if block:
                content.append(block)

        for uploaded in uploaded_examples or []:
            content.append(uploaded_file_to_content_block(uploaded, f"Uploaded example proposal: {uploaded.name}"))

        for uploaded in project_docs or []:
            content.append(uploaded_file_to_content_block(uploaded, f"New project document: {uploaded.name}"))

        if project_notes.strip():
            content.append(text_block("User project notes", project_notes.strip()))

        content.append({"type": "text", "text": PROPOSAL_USER_PROMPT})

        with st.status("Drafting proposal with Claude...", expanded=True) as status:
            st.write("Sending proposal references and project documents to Claude...")
            message = call_claude(client, model, max_tokens, PROPOSAL_SYSTEM_PROMPT, content)
            st.write("Parsing proposal JSON response...")
            raw_text = extract_text(message)
            try:
                context = parse_json_response(raw_text)
                context["proposal_date"] = datetime.now().strftime("%B %d, %Y")

                # Backward-compatible client field cleanup
                if not context.get("client_name") or context.get("client_name") == "Not provided":
                    context["client_name"] = context.get("client_company", "Not provided")

                if not context.get("client_address"):
                    context["client_address"] = "Not provided"

                if not context.get("client_contact"):
                    context["client_contact"] = "Not provided"


            except Exception as exc:
                st.error("Claude returned output that could not be parsed as JSON.")
                st.text_area("Raw Claude output", raw_text, height=400)
                st.exception(exc)
                st.stop()

            st.write("Creating Word proposal draft...")
            docx_bytes = render_proposal_docx(context, terms_region)
            status.update(label="Proposal draft generated", state="complete")

        st.success("Done. Download the proposal draft below.")
        st.download_button(
            label="Download Proposal Draft (.docx)",
            data=docx_bytes,
            file_name="GEOptimize_Proposal_Draft.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        with st.expander("View generated JSON"):
            st.json(context)

            
elif page == "Vertical GHX CAD Request Form":
    st.header("Vertical GHX CAD Request Form")
    st.caption("Use vertical GLD output plus user inputs to prepare a CAD request form.")

    st.subheader("1. Vertical GLD Output PDF")
    cad_gld_pdf = st.file_uploader(
        "Upload vertical GLD output PDF",
        type=["pdf"],
        key="cad_gld_pdf",
    )

    st.subheader("2. User Inputs")

    job_address_name = st.text_input("Job address / name for CAD title block", key="vertical_job")
    timeline_due_date = st.text_input("Timeline / due date", key="vertical_due_date")
    ghx_location_notes = st.text_area("Rough sketch / GHX location notes", height=120, key="vertical_notes")

    rough_sketch_image = st.file_uploader(
        "Upload rough sketch image",
        type=["png", "jpg", "jpeg"],
        key="vertical_rough_sketch_image",
    )

    st.subheader("3. CAD-Specific Inputs")
    st.caption("These are typically not available in GLD and should be confirmed by the designer.")

    col1, col2 = st.columns(2)

    with col1:
        manifold_size = st.text_input("Return / Supply manifold size", key="vertical_manifold_size")
        manifold_material = st.selectbox(
            "Return / Supply manifold material grade",
            ["SDR 11", "SDR 13.5", "Not provided"],
            key="vertical_manifold_material",
        )
        balancing_valve = st.selectbox(
            "Supply manifold balancing valve",
            ["YES", "NO", "Not provided"],
            key="vertical_balancing_valve",
        )
        runout_size = st.text_input("Return / Supply runout size", key="vertical_runout_size")
        runout_material = st.selectbox(
            "Return / Supply runout material grade",
            ["SDR 11", "SDR 13.5", "Not provided"],
            key="vertical_runout_material",
        )
        penetration_type = st.selectbox(
            "Return / Supply penetration",
            ["Wall PEN", "FLOOR PEN", "Not provided"],
            key="vertical_penetration",
        )

    with col2:
        header_size_sequence = st.text_input(
            'Header size sequence, e.g. 3" to 2" to 1-1/4"',
            key="vertical_header_sequence",
        )
        circuits_in_series = st.selectbox(
            "Circuits in series",
            ["NO", "YES", "Not provided"],
            key="vertical_circuits_in_series",
        )
        energy_meter_required = st.selectbox(
            "Energy meter required",
            ["NO", "YES", "Not provided"],
            key="vertical_energy_meter",
        )
        extra_notes = st.text_area("Additional CAD notes", height=120, key="vertical_extra_notes")

    with st.expander("Fields auto-extracted from the Vertical GLD PDF"):
        st.write("The app will ask Claude to extract these fields from the vertical GLD output:")
        st.markdown(
            """
- Borehole number and circuit number
- Borehole length / U-tube depth
- Total bore length
- Borehole spacing and vertical grid arrangement
- Bores per circuit
- U-tube size and material grade from pipe type
- Borehole diameter
- Fluid information
- System flow rates
            """.strip()
        )

    with st.expander("Optional GLD-derived field overrides"):
        st.caption("Leave these blank to let Claude extract them from the vertical GLD output.")
        override_borehole_spacing = st.text_input("Override borehole spacing", key="vertical_override_spacing")
        override_u_tube_size = st.text_input("Override U-tube size", key="vertical_override_u_tube_size")
        override_u_tube_material = st.selectbox(
            "Override U-tube material grade",
            ["", "SDR 11", "SDR 13.5", "Not provided"],
            key="vertical_override_u_tube_material",
        )
        override_u_tube_depth = st.text_input("Override U-tube depth", key="vertical_override_depth")
        override_circuit_number = st.text_input("Override circuit number", key="vertical_override_circuit")

    if st.button("Generate Vertical CAD Request Form", type="primary"):
        if not api_key:
            st.error("Enter your Anthropic API key in the sidebar or .env file.")
            st.stop()

        if not cad_gld_pdf:
            st.error("Upload the vertical GLD output PDF.")
            st.stop()

        client = Anthropic(api_key=api_key)

        cad_prompt = f"""
Analyze the uploaded vertical GLD output PDF and prepare structured information for a Vertical GHX CAD request form.

Use the GLD output to automatically extract these values when available:
- total bore length
- borehole number
- borehole length
- vertical grid arrangement
- borehole separation
- bores per circuit
- pipe type
- pipe material grade, such as SDR11 or SDR13.5
- U-tube size
- U-tube depth
- borehole diameter
- fluid information
- system flow rates

For CAD form fields:
- borehole_spacing should come from Borehole Separation and vertical grid arrangement. Include both if available, e.g. "20.0 ft spacing, 8 x 1 grid".
- u_tube_size should come from Pipe Type.
- u_tube_material should come from Pipe Type or Flow Type if SDR is shown.
- u_tube_depth should come from Borehole Length.
- circuit_number should be inferred from Borehole Number and Bores Per Circuit. For example, 8 boreholes with 1 bore per circuit means 8 circuits.
- extra_notes should summarize useful GLD-derived design context such as total bore length, borehole number, borehole diameter, flow rates, fluid, and grid arrangement.

Use user-entered fields as higher priority than the PDF.
For optional override fields, if the override is blank, use the GLD value.

Return strict JSON only with exactly these keys:
{{
  "job_address_name": "string",
  "timeline_due_date": "string",
  "ghx_location_notes": "string",
  "manifold_size": "string",
  "manifold_material": "string",
  "balancing_valve": "string",
  "runout_size": "string",
  "runout_material": "string",
  "penetration_type": "string",
  "header_size_sequence": "string",
  "borehole_spacing": "string",
  "u_tube_size": "string",
  "u_tube_material": "string",
  "u_tube_depth": "string",
  "circuit_number": "string",
  "circuits_in_series": "string",
  "energy_meter_required": "string",
  "extra_notes": "string"
}}

USER INPUTS:
Job address/name: {job_address_name}
Timeline/due date: {timeline_due_date}
GHX location notes: {ghx_location_notes}
Manifold size: {manifold_size}
Manifold material: {manifold_material}
Balancing valve: {balancing_valve}
Runout size: {runout_size}
Runout material: {runout_material}
Penetration: {penetration_type}
Header size sequence: {header_size_sequence}
Circuits in series: {circuits_in_series}
Energy meter required: {energy_meter_required}
Extra notes: {extra_notes}

OPTIONAL GLD-DERIVED OVERRIDES:
Borehole spacing override: {override_borehole_spacing}
U-tube size override: {override_u_tube_size}
U-tube material override: {override_u_tube_material}
U-tube depth override: {override_u_tube_depth}
Circuit number override: {override_circuit_number}

If a value is not available, return "Not provided".
"""

        content = [
            pdf_block(cad_gld_pdf),
            {"type": "text", "text": cad_prompt},
        ]

        with st.status("Generating Vertical CAD request form...", expanded=True) as status:
            message = call_claude(
                client,
                model,
                max_tokens,
                "You prepare concise structured CAD request information for vertical GHX drawings.",
                content,
            )

            raw_text = extract_text(message)

            try:
                context = parse_json_response(raw_text)
            except Exception as exc:
                st.error("Claude returned output that could not be parsed as JSON.")
                st.text_area("Raw Claude output", raw_text, height=400)
                st.exception(exc)
                st.stop()

            docx_bytes = render_cad_request_docx(context, rough_sketch_image)
            status.update(label="Vertical GHX CAD Request Form", state="complete")

        st.success("Done. Download the Vertical CAD request form below.")
        st.download_button(
            label="Download Vertical CAD Request Form (.docx)",
            data=docx_bytes,
            file_name="Vertical_GHX_CAD_Request_Form.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        with st.expander("View generated JSON"):
            st.json(context)


elif page == "Slinky GHX CAD Request Form":
    st.header("Slinky GHX CAD Request Form")
    st.caption("Use a Slinky GLD output report plus user inputs to prepare a CAD request form.")

    st.subheader("1. Slinky GLD Output PDF")
    slinky_gld_pdf = st.file_uploader(
        "Upload Slinky GLD output PDF",
        type=["pdf"],
        key="slinky_gld_pdf",
    )

    st.subheader("2. User Inputs")

    job_address_name = st.text_input("Job address / name for CAD title block", key="slinky_job")
    timeline_due_date = st.text_input("Timeline / due date", key="slinky_due_date")
    ghx_location_notes = st.text_area("Rough sketch / GHX location notes", height=120, key="slinky_notes")

    rough_sketch_image = st.file_uploader(
        "Upload rough sketch image",
        type=["png", "jpg", "jpeg"],
        key="slinky_rough_sketch_image",
    )

    st.subheader("3. CAD-Specific Inputs")
    st.caption("These are typically not available in GLD and should be confirmed by the designer.")

    col1, col2 = st.columns(2)

    with col1:
        manifold_size = st.text_input("Return / Supply manifold size", key="slinky_manifold_size")
        manifold_material = st.selectbox(
            "Return / Supply manifold material grade",
            ["SDR 11", "SDR 13.5", "Not provided"],
            key="slinky_manifold_material",
        )
        balancing_valve = st.selectbox(
            "Supply manifold balancing valve",
            ["YES", "NO", "Not provided"],
            key="slinky_balancing_valve",
        )
        runout_size = st.text_input("Return / Supply runout size", key="slinky_runout_size")
        runout_material = st.selectbox(
            "Return / Supply runout material grade",
            ["SDR 11", "SDR 13.5", "Not provided"],
            key="slinky_runout_material",
        )
        penetration_type = st.selectbox(
            "Return / Supply penetration",
            ["Wall PEN", "FLOOR PEN", "Not provided"],
            key="slinky_penetration_type",
        )

    with col2:
        header_size_sequence = st.text_input(
            'Header size sequence, e.g. 3" to 2" to 1-1/4"',
            key="slinky_header_size_sequence",
        )
        energy_meter_required = st.selectbox(
            "Energy meter required",
            ["NO", "YES", "Not provided"],
            key="slinky_energy_meter_required",
        )
        extra_notes = st.text_area("Additional CAD notes", height=120, key="slinky_extra_notes")

    with st.expander("Fields auto-extracted from the Slinky GLD PDF"):
        st.write("The app will ask Claude to extract these fields from the GLD output:")
        st.markdown(
            """
- Slinky pipe size and material grade
- Slinky circuit number from trench number
- Slinky overall spacing from trench length, trench count, separation, and total area
- Slinky loop pitch
- Slinky separation
- Trench depth and loop diameter
- Total trench length, single trench length, total pipe length, and total area
- System flow rate
            """.strip()
        )

    if st.button("Generate Slinky CAD Request Form", type="primary"):
        if not api_key:
            st.error("Enter your Anthropic API key in the sidebar or .env file.")
            st.stop()

        if not slinky_gld_pdf:
            st.error("Upload the Slinky GLD output PDF.")
            st.stop()

        client = Anthropic(api_key=api_key)

        slinky_prompt = f"""
Analyze the uploaded Slinky GLD output PDF and prepare structured information for a Slinky GHX CAD request form.

User-entered fields have priority over the PDF. For all other fields, extract values directly from the GLD output PDF.

Extract these values automatically from the GLD output when available:
- Slinky pipe size from Pipe Type, for example 3/4 in. (20 mm)
- Slinky material grade from Pipe Type, for example SDR11 or SDR 11
- Slinky circuit number from Trench Number
- Slinky loop pitch from Loop Pitch
- Slinky separation from Trench Layout Separation
- Trench depth from Trench Layout Depth
- Loop diameter from Loop Diameter
- Total trench length
- Single trench length
- Total pipe length
- Single trench pipe length
- Total area
- System flow rate

For slinky_overall_spacing:
- If the GLD report provides Total Area, Trench Number, Separation, and Single Trench Length, infer the overall spacing as approximately:
  Trench field width = Trench Number x Separation
  Trench field length = heating Single Trench Length if available, otherwise cooling Single Trench Length
- Return as: "approximately [width] ft x [length] ft ([area] ft² total area)"
- If there is not enough information, return the Total Area only.

For extra_notes:
- Include useful GLD-derived CAD notes such as trench depth, loop diameter, total trench length, total pipe length, and system flow rate.
- Preserve any user-entered extra notes and add GLD-derived notes after them.

Return strict JSON only with exactly these keys:
{{
  "job_address_name": "string",
  "timeline_due_date": "string",
  "ghx_location_notes": "string",
  "manifold_size": "string",
  "manifold_material": "string",
  "balancing_valve": "string",
  "runout_size": "string",
  "runout_material": "string",
  "penetration_type": "string",
  "header_size_sequence": "string",
  "slinky_pipe_size": "string",
  "slinky_material": "string",
  "slinky_overall_spacing": "string",
  "slinky_circuit_number": "string",
  "slinky_loop_pitch": "string",
  "slinky_separation": "string",
  "energy_meter_required": "string",
  "extra_notes": "string"
}}

USER INPUTS:
Job address/name: {job_address_name}
Timeline/due date: {timeline_due_date}
GHX location notes: {ghx_location_notes}
Manifold size: {manifold_size}
Manifold material: {manifold_material}
Balancing valve: {balancing_valve}
Runout size: {runout_size}
Runout material: {runout_material}
Penetration: {penetration_type}
Header size sequence: {header_size_sequence}
Energy meter required: {energy_meter_required}
Extra notes: {extra_notes}

If a user-entered value is blank, use "Not provided" unless the value can be extracted from the GLD output.
If a GLD-derived value is not available, return "Not provided".
Do not invent values.
"""

        content = [
            pdf_block(slinky_gld_pdf),
            {"type": "text", "text": slinky_prompt},
        ]

        with st.status("Generating Slinky CAD request form...", expanded=True) as status:
            message = call_claude(
                client,
                model,
                max_tokens,
                "You prepare concise structured CAD request information for slinky GHX drawings.",
                content,
            )

            raw_text = extract_text(message)

            try:
                context = parse_json_response(raw_text)
            except Exception as exc:
                st.error("Claude returned output that could not be parsed as JSON.")
                st.text_area("Raw Claude output", raw_text, height=400)
                st.exception(exc)
                st.stop()

            docx_bytes = render_slinky_cad_docx(context, rough_sketch_image)
            status.update(label="Slinky GHX CAD Request Form", state="complete")

        st.success("Done. Download the Slinky CAD request form below.")
        st.download_button(
            label="Download Slinky CAD Request Form (.docx)",
            data=docx_bytes,
            file_name="Slinky_GHX_CAD_Request_Form.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        with st.expander("View generated JSON"):
            st.json(context)


elif page == "HDD GHX CAD Request Form":
    st.header("Horizontal Directionally Drilled GHX CAD Request Form")
    st.caption("Use an HDD GLD output report plus user inputs to prepare a CAD request form.")

    st.subheader("1. HDD GLD Output PDF")
    hdd_gld_pdf = st.file_uploader(
        "Upload HDD GLD output PDF",
        type=["pdf"],
        key="hdd_gld_pdf",
    )

    st.subheader("2. User Inputs")

    job_address_name = st.text_input("Job address / name for CAD title block", key="hdd_job")
    timeline_due_date = st.text_input("Timeline / due date", key="hdd_due_date")
    ghx_location_notes = st.text_area("Rough sketch / GHX location notes", height=120, key="hdd_notes")

    rough_sketch_image = st.file_uploader(
        "Upload rough sketch image",
        type=["png", "jpg", "jpeg"],
        key="hdd_rough_sketch_image",
    )

    st.subheader("3. CAD-Specific Inputs")
    st.caption("These are typically not available in GLD and should be confirmed by the designer.")

    col1, col2 = st.columns(2)

    with col1:
        manifold_size = st.text_input("Return / Supply manifold size", key="hdd_manifold_size")
        manifold_material = st.selectbox(
            "Return / Supply manifold material grade",
            ["SDR 11", "SDR 13.5", "Not provided"],
            key="hdd_manifold_material",
        )
        balancing_valve = st.selectbox(
            "Supply manifold balancing valve",
            ["YES", "NO", "Not provided"],
            key="hdd_balancing_valve",
        )
        runout_size = st.text_input("Return / Supply runout size", key="hdd_runout_size")
        runout_material = st.selectbox(
            "Return / Supply runout material grade",
            ["SDR 11", "SDR 13.5", "Not provided"],
            key="hdd_runout_material",
        )
        penetration_type = st.selectbox(
            "Return / Supply penetration",
            ["Wall PEN", "FLOOR PEN", "Not provided"],
            key="hdd_penetration",
        )

    with col2:
        header_size_sequence = st.text_input(
            'Header size sequence, e.g. 3" to 2" to 1-1/4"',
            key="hdd_header_sequence",
        )
        energy_meter_required = st.selectbox(
            "Energy meter required",
            ["NO", "YES", "Not provided"],
            key="hdd_energy_meter",
        )
        extra_notes = st.text_area("Additional CAD notes", height=120, key="hdd_extra_notes")

    with st.expander("Fields auto-extracted from the HDD GLD PDF"):
        st.write("The app will ask Claude to extract these fields from the GLD output:")
        st.markdown(
            """
- Horizontal pipe size and material grade
- Horizontal circuit number from trench number
- Horizontal overall spacing from trench length, trench count, separation, and total area
- Number of horizontal layers from Pipe Layout [X x Y]
- Depth of each horizontal layer from trench depth and vertical separation
- Trench number, single trench length, total trench length, total pipe length, and total area
- Horizontal and vertical pipe separation
            """.strip()
        )

    if st.button("Generate HDD CAD Request Form", type="primary"):
        if not api_key:
            st.error("Enter your Anthropic API key in the sidebar or .env file.")
            st.stop()

        if not hdd_gld_pdf:
            st.error("Upload the HDD GLD output PDF.")
            st.stop()

        client = Anthropic(api_key=api_key)

        hdd_prompt = f"""
Analyze the uploaded horizontal directionally drilled GLD output PDF and prepare structured information for an HDD GHX CAD request form.

Use the GLD output to extract values wherever available. Use user-entered values as higher priority than the PDF.

The HDD GLD report may use horizontal/trench terminology. Map it to the CAD form as follows:
- Horizontal pipe size: extract from Pipe Type, for example 1 1/4 in. (32 mm)
- Horizontal material grade: extract from Pipe Type, for example SDR11 or SDR 11
- Horizontal circuit number: use Trench Number
- Horizontal overall spacing: summarize approximate layout using single trench length, trench number, trench separation, and total area. If possible, format as "approx. length x width".
- How many horizontal layers: infer from Pipe Layout [X x Y]. If the layout is 2 x 2, return 2 horizontal layers.
- Depth of each horizontal layer: use Trench Layout Depth and Vertical Separation. If depth is 20.0 ft and vertical separation is 120.0 in, return "Layer 1: 20 ft, Layer 2: 30 ft".
- Include trench separation, trench depth, horizontal separation, vertical separation, total trench length, single trench length, total pipe length, and total area in extra_notes if useful.

Return strict JSON only with exactly these keys:
{{
  "job_address_name": "string",
  "timeline_due_date": "string",
  "ghx_location_notes": "string",
  "manifold_size": "string",
  "manifold_material": "string",
  "balancing_valve": "string",
  "runout_size": "string",
  "runout_material": "string",
  "penetration_type": "string",
  "header_size_sequence": "string",
  "horizontal_pipe_size": "string",
  "horizontal_material": "string",
  "horizontal_overall_spacing": "string",
  "horizontal_circuit_number": "string",
  "horizontal_layers": "string",
  "horizontal_layer_depths": "string",
  "energy_meter_required": "string",
  "extra_notes": "string"
}}

USER INPUTS:
Job address/name: {job_address_name}
Timeline/due date: {timeline_due_date}
GHX location notes: {ghx_location_notes}
Manifold size: {manifold_size}
Manifold material: {manifold_material}
Balancing valve: {balancing_valve}
Runout size: {runout_size}
Runout material: {runout_material}
Penetration: {penetration_type}
Header size sequence: {header_size_sequence}
Energy meter required: {energy_meter_required}
Extra notes: {extra_notes}

If a value is missing, return "Not provided".
Do not invent values.
"""

        content = [
            pdf_block(hdd_gld_pdf),
            {"type": "text", "text": hdd_prompt},
        ]

        with st.status("Generating HDD CAD request form...", expanded=True) as status:
            message = call_claude(
                client,
                model,
                max_tokens,
                "You prepare concise structured CAD request information for horizontal directionally drilled GHX drawings.",
                content,
            )

            raw_text = extract_text(message)

            try:
                context = parse_json_response(raw_text)
            except Exception as exc:
                st.error("Claude returned output that could not be parsed as JSON.")
                st.text_area("Raw Claude output", raw_text, height=400)
                st.exception(exc)
                st.stop()

            docx_bytes = render_hdd_cad_docx(context, rough_sketch_image)
            status.update(label="HDD GHX CAD Request Form", state="complete")

        st.success("Done. Download the HDD CAD request form below.")
        st.download_button(
            label="Download HDD CAD Request Form (.docx)",
            data=docx_bytes,
            file_name="HDD_GHX_CAD_Request_Form.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        with st.expander("View generated JSON"):
            st.json(context)
